
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_docx_from_md(md_path, docx_path):
    document = Document()
    
    # Simple Markdown Parser for Headers and Tables
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_buffer = []
    
    for line in lines:
        line = line.strip()
        
        # Handle Table Rows
        if line.startswith('|'):
            table_buffer.append(line)
            continue
        else:
            # Process buffered table if any
            if table_buffer:
                process_table(document, table_buffer)
                table_buffer = []

        if not line:
            continue

        # Handle Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            document.add_heading(text, level=level)
        
        # Handle Blockquotes
        elif line.startswith('>'):
            p = document.add_paragraph(line.lstrip('>').strip())
            p.italic = True
        
        # Handle Lists
        elif line.startswith('* ') or line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        
        # Handle Standard Paragraphs
        else:
            document.add_paragraph(line)

    # Process any remaining table at the end
    if table_buffer:
        process_table(document, table_buffer)

    document.save(docx_path)
    print(f"Successfully created {docx_path}")

def process_table(document, table_lines):
    # Filter out separator lines (e.g., |---|---|)
    rows = [row for row in table_lines if '---' not in row]
    
    if not rows:
        return

    # Determine dimensions
    first_row_cells = rows[0].strip('|').split('|')
    num_cols = len(first_row_cells)
    
    table = document.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_content in enumerate(rows):
        cells = row_content.strip('|').split('|')
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(cells):
            if j < len(row_cells):
                row_cells[j].text = cell_text.strip().replace('**', '')

if __name__ == "__main__":
    md_file = "ndcg.md"
    docx_file = "ndcg.docx"
    
    if os.path.exists(md_file):
        create_docx_from_md(md_file, docx_file)
    else:
        print(f"Error: {md_file} not found.")
